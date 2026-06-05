#!/usr/bin/env python3
"""Convert a clean Markdown resume draft into a polished Word DOCX."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_border(paragraph, color: str = "D9E2EC") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_formatted_runs(paragraph, text: str, *, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = bold_default
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        run = paragraph.add_run(part)
        set_east_asia_font(run, "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, bold, color in [
        ("Heading 1", 20, True, "111827"),
        ("Heading 2", 11.5, True, "0F4C81"),
        ("Heading 3", 10.5, True, "111827"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(6 if style_name != "Heading 1" else 0)
        style.paragraph_format.space_after = Pt(3)


def clean_line(line: str) -> str:
    return line.strip().replace("｜", " | ")


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        paragraph = doc.add_paragraph(style="Heading 1")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_east_asia_font(run, "Microsoft YaHei")
        run.bold = True
        run.font.size = Pt(20)
        return

    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    add_formatted_runs(paragraph, text, bold_default=True)
    if level == 2:
        set_paragraph_border(paragraph)


def add_body_paragraph(doc: Document, text: str, *, centered: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    add_formatted_runs(paragraph, text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(2)
    add_formatted_runs(paragraph, text)


def markdown_to_docx(markdown: str, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    saw_title = False
    previous_was_title = False

    for raw_line in markdown.splitlines():
        line = clean_line(raw_line)
        if not line or line == "---":
            continue
        if line.startswith(">"):
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_heading(doc, text, level)
            saw_title = saw_title or level == 1
            previous_was_title = level == 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            add_bullet(doc, bullet_match.group(1).strip())
            previous_was_title = False
            continue

        add_body_paragraph(doc, line, centered=(saw_title and previous_was_title))
        previous_was_title = False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert a clean Markdown resume into DOCX.")
    parser.add_argument("input", help="Input Markdown resume path.")
    parser.add_argument("output", help="Output DOCX path.")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f"Input file does not exist: {input_path}", file=sys.stderr)
        return 2

    markdown = input_path.read_text(encoding="utf-8")
    markdown_to_docx(markdown, output_path)

    check_doc = Document(output_path)
    paragraph_count = len([p for p in check_doc.paragraphs if p.text.strip()])
    if paragraph_count == 0:
        print("DOCX validation failed: no non-empty paragraphs found.", file=sys.stderr)
        return 1

    print(f"Wrote {output_path} ({paragraph_count} non-empty paragraphs).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
